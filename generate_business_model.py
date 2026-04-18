from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Colours ───────────────────────────────────────────────────────────────────
BLACK    = RGBColor(0x1D, 0x1D, 0x1F)
BLUE     = RGBColor(0x00, 0x66, 0xCC)
MUTED    = RGBColor(0x6E, 0x6E, 0x73)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
MID_GREY = RGBColor(0xD1, 0xD1, 0xD6)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_border(table, color='D1D1D6', size=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_top_border(para, color='0066CC', size=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)

def h1(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(16)
    run.font.bold      = True
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after  = Pt(6)
    add_top_border(p)
    return p

def h2(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(12)
    run.font.bold      = True
    run.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.bold      = True
    run.font.color.rgb = MUTED
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, italic=False, color=None, size=10.5):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.italic    = italic
    run.font.color.rgb = color or BLACK
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    return p

def bullet(text, indent=False):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(t)
    # header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, '1D1D1F')
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.font.bold      = True
        run.font.size      = Pt(9.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
    # rows
    for ri, row_data in enumerate(rows):
        row  = t.rows[ri + 1]
        fill = 'F5F5F7' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size      = Pt(9.5)
            run.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def callout(text):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F7')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size      = Pt(9.5)
    run.font.color.rgb = BLACK
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('3lleven')
r.font.size = Pt(48); r.font.bold = True; r.font.color.rgb = BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('Business Model & Day One Blueprint')
r.font.size = Pt(16); r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Web Design & Development — Built for Scale')
r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('─' * 62)
r.font.color.rgb = MID_GREY; r.font.size = Pt(9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Version 2.0  ·  April 2026  ·  Confidential')
r.font.size = Pt(10); r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. THE BRAND
# ══════════════════════════════════════════════════════════════════════════════
h1('1. The Brand')

h2('The Story')
body(
    '3lleven is named after the founder\'s birthday — March 11th (3/11). The name represents a personal '
    'commitment to the work and signals a "new beginning" for every business 3lleven works with. '
    'The "3" functions as leetspeak for the letter "E," quietly signaling a coding background to anyone '
    'who recognizes it — while still reading as a clean, memorable brand name to everyone else.'
)

h2('The Look')
body(
    'High-tech, clean, and professional. The aesthetic draws from premium product design — think less '
    '"freelancer portfolio," more "software company." Every deliverable, from proposals to the finished '
    'site, should reinforce the same standard.'
)

h2('Domain Strategy')
table(
    headers = ['Domain', 'Purpose', 'Priority'],
    rows    = [
        ['3lleven.dev',      'Primary — signals technical credibility to other devs and businesses', 'First choice'],
        ['3lleven.studio',   'Alternative — creative-agency feel if .dev is taken', 'Second choice'],
        ['threeeleven.com',  'Redirect only — captures typo traffic from people who can\'t remember the "3" spelling', 'Buy regardless'],
    ],
    widths = [1.8, 3.8, 1.6],
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. SERVICE & PRICING MODEL
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Service & Pricing Model')
body(
    '3lleven operates on value-based pricing — not hourly rates. Every package is fixed-fee so clients '
    'know exactly what they\'re paying before work begins. There are two revenue streams: one-time builds '
    '(the "Engine") and recurring monthly plans (the "Fuel").',
    color=MUTED, italic=True
)

h2('A. One-Time Builds — The "Engine" Build')
body('A fixed-fee, fully custom site. The client gets the running website. Source code transfer is a separate fee (see Section 5).')
doc.add_paragraph()

table(
    headers = ['Item', 'Price', 'Notes'],
    rows    = [
        ['New Build Base',              '$750',    '1–3 page site. Custom design, mobile responsive, light/dark mode, hover effects, scroll animations, contact form, basic SEO, SSL, domain & hosting setup, 2 revisions.'],
        ['Redesign Base',               '$300',    'Client already has a site — refresh the design, keep the domain. All add-ons apply at same price.'],
        ['Additional Pages',            '$75 each','Per page beyond the included 3.'],
        ['Advanced Animations',         '+$150',   'Framer Motion or GSAP — parallax, staggered sequences, page transitions. Beyond the standard scroll fades included in every build.'],
        ['E-Commerce Integration',      '+$300',   'Stripe or Shopify — product catalog, cart, checkout, payment processing.'],
        ['Blog / CMS',                  '+$100',   'Sanity.io — client can update content without touching code.'],
        ['Booking / Appointment System','+$150',   'Integrated scheduling for services, consultations, or appointments.'],
        ['Full SEO Package',            '+$100',   'Keyword research, schema markup, Google Search Console, sitemap.'],
        ['Client Portal / Dashboard',   '+$150',   'Password-protected area for client files, invoices, or account data.'],
        ['API / Third-Party Integration','+$100',  'CRM, email marketing, inventory, or any external service connection.'],
    ],
    widths = [2.0, 1.4, 3.9],
)

h3('Payment Terms')
for b in [
    '50% deposit collected via Stripe before any work begins',
    '50% collected before the site goes live — not after, not on a promise',
    'For projects over $1,000: 50% deposit / 25% at design approval / 25% at launch',
    'Late payment: +5% fee after 7 days past due (stated in every contract)',
]:
    bullet(b)

h2('B. Monthly Plans — The "Fuel" / Recurring Revenue')
body('Three tiers. Each client is onboarded to a plan at launch. This is the business model that compounds over time.')
doc.add_paragraph()

table(
    headers = ['Tier', 'Price', 'Target Client', 'Key Value'],
    rows    = [
        ['Tier 1 — Basic',   '$199/mo',      '1-Page "Digital Business Card"', 'Hosting, security, backups, and peace of mind. Set it and forget it.'],
        ['Tier 2 — Growth',  '$299–$399/mo', 'Standard Business (5+ pages)',   'SEO monitoring, monthly content updates, analytics reporting.'],
        ['Tier 3 — Revenue', '$499–$699/mo', 'E-Commerce / Custom Apps',        'Daily backups, priority support, transactional security, advanced monitoring.'],
    ],
    widths = [1.5, 1.4, 2.0, 2.4],
)

h3('What Each Tier Includes')
body('Tier 1 — Basic ($199/mo)', italic=False, color=BLUE)
for b in ['Hosting & SSL management','Security & dependency updates','Monthly automated backups','24/7 uptime monitoring','1 hr dev support per month','Monthly performance report']:
    bullet(b)
doc.add_paragraph()
body('Tier 2 — Growth ($299–$399/mo)', italic=False, color=BLUE)
for b in ['Everything in Tier 1','5 hrs dev/design work per month (rolls over up to 2 months)','Monthly SEO report & recommendations','Google Business profile management','2 content updates per month','Priority next-business-day response','$299 for standard sites · $399 for larger or CMS-heavy sites']:
    bullet(b)
doc.add_paragraph()
body('Tier 3 — Revenue ($499–$699/mo)', italic=False, color=BLUE)
for b in ['Everything in Tier 2','E-commerce management & optimization','Daily automated backups','Up to 10 hrs dev work per month','Priority 24-hour support','Unlimited content updates','Advanced analytics & conversion tracking','Custom integrations maintenance','$499 standard e-commerce · $699+ large or complex stores']:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Tech Stack — Your Technical Edge')
body(
    'As a CS graduate, 3lleven\'s advantage over Wix/Squarespace resellers and template-flippers is '
    'engineering. The stack below produces sites that score 100 on Lighthouse — and that score '
    'is the primary sales tool against "drag-and-drop" competitors.'
)

table(
    headers = ['Layer', 'Tool', 'Why'],
    rows    = [
        ['Build Framework',   'Astro or Next.js',  'Astro for static/content sites (zero JS by default = perfect Lighthouse score). Next.js for dynamic apps, e-commerce, or anything needing a backend.'],
        ['Styling',           'Tailwind CSS',       'Rapid, consistent, professional UI. No bloated CSS files. Easy to hand off or maintain.'],
        ['CMS',               'Sanity.io',          'Structured content — clients edit text and images through a clean dashboard without ever touching your code or breaking the layout.'],
        ['Deployment',        'Vercel or Netlify',  'Connected to GitHub. Every push to main = automatic live deployment. Preview URLs for every branch.'],
        ['Payments (Client)', 'Stripe',             'For client sites needing e-commerce or bookings. Handles checkout, refunds, and subscriptions.'],
        ['Billing (Your Business)', 'Stripe',       'Automated recurring subscriptions for monthly plans. Card fails → service pauses automatically. You never chase a payment.'],
        ['Version Control',   'GitHub',             'Every project has its own private repo. Branching: dev for active work, main for live deployment.'],
    ],
    widths = [1.6, 1.7, 4.0],
)

h2('The Lighthouse Pitch')
body(
    'Run PageSpeed Insights on any competitor\'s Wix or GoDaddy site. Screenshot the score (usually 40–60). '
    'Then show your site scoring 95–100. That gap is your sales pitch in one image. No client needs to '
    'understand web performance — they just need to see the number.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. OPERATIONAL WORKFLOW — THE 11-STEP PLAYBOOK
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Operational Workflow — The Playbook')
body(
    'Follow this exact sequence on every single project. Deviation is how scope creep, '
    'non-payment, and bad clients happen.',
    color=MUTED, italic=True
)

steps = [
    ('Step 1 — Discovery (15 min)', [
        'Short call — not a sales pitch, a qualification call',
        'Key questions: What do you need? What\'s your timeline? What\'s your budget?',
        'If they\'re not a fit (budget too low, unclear goals, bad vibes) — pass politely',
        'If they are a fit — send a Calendly link for the full proposal call',
    ]),
    ('Step 2 — Proposal & Contract + 50% Deposit', [
        'Send a 1-page proposal: scope (exact pages/features), price, timeline, payment terms',
        'Send the contract alongside the proposal — both must be signed before deposit',
        'Collect 50% deposit via Stripe Payment Link',
        'Do not open Figma until the deposit clears. No exceptions.',
    ]),
    ('Step 3 — Figma Design', [
        'Build a low-fidelity wireframe first — layout and structure, no colors or fonts',
        'Get written approval on structure before moving to high-fidelity',
        'Share the Figma link — client comments directly in the file',
        'Max 2 revision rounds. Anything beyond is a change order.',
    ]),
    ('Step 4 — Development', [
        'Build on a private GitHub repo: 3lleven-clientname',
        'Work on the dev branch — never commit directly to main',
        'Use the appropriate stack for the project tier (Astro/Next.js + Tailwind + Sanity)',
        'Target Lighthouse score: 95+ on Performance, Accessibility, SEO',
    ]),
    ('Step 5 — Review (Preview Link)', [
        'Deploy to Vercel/Netlify preview URL',
        'Send the link with a Loom walkthrough video — 3–5 minutes, narrated',
        'Client reviews and responds with feedback — max 2 rounds',
        'Get written "Approved — go live" before touching the domain',
    ]),
    ('Step 6 — Final Payment', [
        'Send the remaining 50% invoice via Stripe',
        'Do not connect the live domain until payment clears',
        'If they stall: the site sits on the preview URL until they pay',
    ]),
    ('Step 7 — Deployment', [
        'Connect the custom domain in Vercel/Netlify',
        'Update DNS records: CNAME www → Vercel, A record @ → Vercel IP',
        'SSL auto-provisions — verify everything live: all pages, form, mobile, speed',
        'Merge dev → main on GitHub to trigger the live deployment',
        'Run final Lighthouse audit and screenshot the score',
    ]),
    ('Step 8 — Onboarding to Monthly Plan', [
        'Every client should leave with a monthly plan — it\'s not optional, it\'s part of the pitch',
        'Frame it as: "Here\'s how I keep your site secure, fast, and updated going forward"',
        'Set up Stripe Subscription for their tier — they receive an automatic invoice each month',
        'If card fails: Stripe retries 3x then the plan pauses. You don\'t chase anyone.',
        'Send a handoff doc: how to update content, logins (via 1Password), your contact info',
    ]),
]

for title, bullets_list in steps:
    h2(title)
    for b in bullets_list:
        bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 5. NO-BS BUSINESS RULES
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Non-Negotiable Business Rules')

h2('Source Code is Not Free')
body(
    'If a client buys a $750 site, they get the live, running website. They do not automatically '
    'receive the source code to take to another developer. The code is the business asset that '
    'keeps them on a monthly plan.'
)
for b in [
    'Source code transfer = separate "Transfer Fee": $100 for new builds, $75 for redesigns',
    'Exception: clients who have been on an active monthly plan for 12+ consecutive months receive the source code at no extra charge as a loyalty reward',
    'This should be clearly stated in every contract — no surprises, no arguments',
]:
    bullet(b)

h2('Automate All Billing — Never Ask Manually')
body(
    'The moment you start manually sending "Hey, just a reminder about your $199 this month" messages, '
    'you\'ve lost. Set up Stripe Subscriptions from day one.'
)
for b in [
    'Stripe sends the invoice automatically on the billing date',
    'If the card fails: Stripe retries 3 times over 7 days, then the subscription pauses',
    'When paused: send one email — "Your plan is paused — update your card here: [link]"',
    'If they don\'t respond in 14 days: the plan cancels. No manual follow-up, no chasing.',
]:
    bullet(b)

h2('The Lighthouse Pitch (Biggest Selling Point)')
body(
    'This is the technical advantage that separates 3lleven from every Fiverr reseller and '
    'GoDaddy website builder. Use it in every pitch.'
)
callout(
    'Pitch script:\n\n'
    '"Let me show you something. Here\'s your current site on Google PageSpeed Insights — '
    'it scores a 43 out of 100. That means Google is actively ranking you lower because of '
    'your website speed. Here\'s what I built for a similar client — 97 out of 100. That '
    'difference shows up in search results and in how fast your site loads on someone\'s phone. '
    'That\'s what you\'re paying for."'
)

h2('Standard Features on Every Build')
body('These are never an upsell. They are the baseline that justifies the $750 starting price.')
for b in [
    'Light & dark mode toggle',
    'Hover & interaction effects on all interactive elements',
    'Scroll-triggered animations (fade in, slide up)',
    'Smooth scroll and page transitions',
    'Mobile responsive (tested on iOS and Android)',
    'Contact form with spam protection',
    'Basic on-page SEO (meta tags, Open Graph, sitemap)',
    'SSL certificate (auto via Vercel/Netlify)',
    'Lighthouse score target: 95+',
]:
    bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# 6. CLIENT ACQUISITION
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Client Acquisition')

h2('Fastest Leads — Start Here')
for b in [
    'Google "[business type] [your city]" — open every site, screenshot the bad ones, cold email the owner',
    'Google Maps — businesses that appear in local results but have no website linked are high-intent leads',
    'Text 20 people personally: "I build websites starting at $750. Know anyone who needs one?"',
    'Visit local businesses in person — a printed one-pager with a QR code to 3lleven.dev works',
]:
    bullet(b)

h2('Cold Outreach Template')
callout(
    'Subject: Your website — quick thought\n\n'
    'Hey [Name],\n\n'
    'I came across [Business Name] and noticed your site is loading slowly on mobile — '
    'that\'s costing you customers who bounce before the page even finishes loading.\n\n'
    'I\'m Adebola, I run 3lleven — a web dev studio that builds custom sites starting at $750. '
    'I can show you exactly what a performance score comparison looks like between your current '
    'site and what I\'d build.\n\n'
    'Worth a 15-minute call?\n\n'
    '— Adebola\n'
    '3lleven.dev'
)

h2('Channels')
table(
    headers = ['Channel', 'Approach', 'Priority'],
    rows    = [
        ['Instagram / TikTok',    '"Before & after" redesign videos. Show the Lighthouse score jump. Short-form is the best organic channel for this niche.', 'High'],
        ['LinkedIn',              'Post case studies, Lighthouse scores, and client wins. Connect with local business owners directly.', 'Medium'],
        ['Facebook Local Groups', 'Join local business owner groups. Provide value for 2–3 weeks, then pitch.', 'Medium'],
        ['Google Maps scraping',  'Find businesses with no website or a broken one. Cold email or DM.', 'High — high intent'],
        ['Referral Program',      '$50 cash or credit for every signed referral. Tell every client explicitly.', 'High — compounds'],
        ['Contra / Upwork',       'Early on, to build reviews. Lower margin but good for portfolio proof.', 'Low — start only'],
    ],
    widths = [1.6, 3.6, 1.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. TOOLS & OPERATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Tools & Operations')

table(
    headers = ['Purpose', 'Tool', 'Cost'],
    rows    = [
        ['Scheduling',              'Calendly',                          'Free'],
        ['Contracts & Invoicing',   'Bonsai',                            '$25/mo — contracts, invoices, time tracking in one place'],
        ['Recurring Billing',       'Stripe Subscriptions',              '2.9% + 30¢ per transaction'],
        ['Design & Mockups',        'Figma',                             'Free'],
        ['Client Walkthroughs',     'Loom',                              'Free (5 min) / $15/mo unlimited'],
        ['File Sharing',            'Google Drive',                      'Free'],
        ['Project Tracking',        'Notion',                            'Free'],
        ['Business Email',          'Google Workspace (yourname@3lleven.dev)', '$6/mo'],
        ['Uptime Monitoring',       'UptimeRobot',                       'Free — monitors up to 50 sites'],
        ['Domains',                 'Namecheap or Porkbun',              '~$12/yr per domain'],
        ['Version Control',         'GitHub (private repos)',            'Free'],
        ['Credential Sharing',      'Bitwarden or 1Password',            'Free / $3/mo'],
        ['Performance Testing',     'Google PageSpeed Insights',         'Free — use in every pitch'],
    ],
    widths = [1.8, 2.5, 3.0],
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. REVENUE PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Revenue Projections')

table(
    headers = ['Milestone', 'Projects/mo', 'Avg Project', 'Monthly Recurring', 'Est. Monthly Total'],
    rows    = [
        ['Month 1–2 (Starting)',   '1–2',  '$900',    '$0',     '~$900–$1,800'],
        ['Month 3 (3 months in)',  '2–3',  '$1,000',  '$400',   '~$2,400'],
        ['Month 6',                '3–4',  '$1,100',  '$1,000', '~$4,300'],
        ['Month 12',               '4–5',  '$1,200',  '$2,500', '~$7,300'],
        ['Month 18',               '4–5',  '$1,400',  '$5,000', '~$10,000'],
    ],
    widths = [1.8, 1.3, 1.4, 1.8, 1.9],
)

body(
    'The recurring revenue line is the most important number. Every monthly plan client compounds — '
    'they pay every month whether you\'re building or not. Prioritize converting every project client '
    'into a monthly plan at launch. A $750 client on Tier 1 ($199/mo) is worth $2,388 in year one alone.',
    color=MUTED, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('─' * 62)
r.font.color.rgb = MID_GREY; r.font.size = Pt(9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('3lleven  ·  3lleven.dev  ·  Confidential  ·  Version 2.0  ·  April 2026')
r.font.size = Pt(8.5); r.font.color.rgb = MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\d3bol\Documents\3lleven_Business_Model.docx'
doc.save(out)
print(f'Saved: {out}')
